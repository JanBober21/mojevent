from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth import login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import AuthenticationForm
from django.contrib import messages
from django.db.models import Q, Avg, Count, Max
from django.http import JsonResponse
from django.utils import timezone
from datetime import datetime, timedelta, date
import json
import math

from .models import Restaurant, Booking, Review, RestaurantOwner, BookingNote, BookingTodo, Menu, MenuItem, BookingMenuItem, AttractionItem, BookingMessage, RestaurantImage, SavedMenu, MenuItemTemplate, BlockedDate, BookingCourse, MenuTypeChoice, Dish
from .forms import BookingForm, ReviewForm, UserRegisterForm, RestaurantSearchForm, OwnerRegisterForm, RestaurantForm, UserSettingsForm
from .style_scraper import scrape_styles
from .social_scraper import import_from_urls, detect_platform, PLATFORM_LABELS, PLATFORM_ICONS
from .calendar_utils import build_month_grid, get_day_details, MONTH_NAMES_PL


# ── Stałe opisujące typy menu ──────────────────────────────────────────────
MENU_TYPE_INFO = [
    {
        "key": "detailed",
        "label": "Szczegółowy plan",
        "icon": "bi-list-ol",
        "color": "primary",
        "desc_owner": (
            "Klient wybiera z pełnej listy pozycje menu i określa dokładną "
            "ilość każdego dania. Idealny gdy chcesz znać precyzyjne "
            "zamówienie przed imprezą."
        ),
        "desc_client": (
            "Wybierz z pełnej listy dania i wpisz ile porcji potrzebujesz. "
            "Widzisz ceny i sumujesz dokładny koszt."
        ),
    },
    {
        "key": "limited",
        "label": "Ograniczony Wybór",
        "icon": "bi-card-checklist",
        "color": "success",
        "desc_owner": (
            "Klient wybiera 2\u20133 dania z każdej kategorii do karty. "
            "Goście na imprezie zamawiają od kelnera po 1 pozycji z karty. "
            "Świetny kompromis między kontrolą a wyborem."
        ),
        "desc_client": (
            "Wybierz 2\u20133 dań z każdej kategorii \u2014 na imprezie "
            "goście sami zdecydują, które danie zamawiają od kelnera."
        ),
    },
    {
        "key": "buffet",
        "label": "Szwedzki Stół",
        "icon": "bi-table",
        "color": "warning",
        "desc_owner": (
            "Klient wybiera dania i ustala kolejność serwowania na stół. "
            "Goście sami sobie nakładają. Idealny na większe imprezy "
            "i przyjęcia."
        ),
        "desc_client": (
            "Wybierz dania i ustal kolejność, w jakiej trafią na stół \u2014 "
            "goście nakładają sami ile chcą."
        ),
    },
    {
        "key": "custom_mix",
        "label": "Własny Mix",
        "icon": "bi-sliders",
        "color": "info",
        "desc_owner": (
            "Klient układa własny plan: definiuje etapy (zupa, drugie danie, "
            "deser\u2026), sposób serwowania i przypisuje dania do grup gości "
            "(dorośli, dzieci, vege). Daje pełną kontrolę."
        ),
        "desc_client": (
            "Ułóż swój plan: dodawaj etapy posiłku (zupa, główne, deser\u2026), "
            "wybieraj dania i określ kto co dostaje."
        ),
    },
]

MENU_TYPE_MAP = {t["key"]: t for t in MENU_TYPE_INFO}


def _get_owner_context(request):
    """Return (memberships_qs, active_restaurant, active_membership) for the current user.

    Uses session key 'active_restaurant_id' to remember which firm is selected.
    Returns (None, None, None) when the user has no firm memberships.
    """
    memberships = RestaurantOwner.objects.filter(
        user=request.user, restaurant__isnull=False,
    ).select_related("restaurant")

    if not memberships.exists():
        # User has an owner record with restaurant=None (just registered)
        if RestaurantOwner.objects.filter(user=request.user).exists():
            return memberships, None, None
        return None, None, None

    active_id = request.session.get("active_restaurant_id")
    active_membership = None

    if active_id:
        active_membership = memberships.filter(restaurant_id=active_id).first()

    if not active_membership:
        active_membership = memberships.first()
        request.session["active_restaurant_id"] = active_membership.restaurant_id

    return memberships, active_membership.restaurant, active_membership


def _haversine_km(lat1, lon1, lat2, lon2):
    """Oblicz odległość w km między dwoma punktami GPS (Haversine)."""
    R = 6371  # promień Ziemi w km
    d_lat = math.radians(lat2 - lat1)
    d_lon = math.radians(lon2 - lon1)
    a = (math.sin(d_lat / 2) ** 2 +
         math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) *
         math.sin(d_lon / 2) ** 2)
    return R * 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))


# ── Strona główna ──────────────────────────────────────────────────────────────

def home(request):
    """Strona główna z wyróżnionymi firmami."""
    featured = Restaurant.objects.filter(is_active=True).annotate(
        avg_rating=Avg("reviews__rating")
    ).order_by("-avg_rating")[:6]

    # Build map markers JSON for firms with GPS coordinates
    all_restaurants = Restaurant.objects.filter(is_active=True)
    markers = []
    for r in all_restaurants:
        if r.latitude and r.longitude:
            markers.append({
                "name": r.name,
                "city": r.city,
                "lat": float(r.latitude),
                "lng": float(r.longitude),
                "price": str(r.price_per_person),
                "max_guests": r.max_guests,
                "pk": r.pk,
                "firm_type": r.get_firm_type_display(),
            })

    return render(request, "bookings/home.html", {
        "featured_restaurants": featured,
        "map_markers_json": json.dumps(markers),
        "restaurant_count": Restaurant.objects.filter(is_active=True).count(),
        "booking_count": Booking.objects.count(),
        "review_count": Review.objects.count(),
    })


# ── Restauracje ────────────────────────────────────────────────────────────────

def restaurant_list(request):
    """Lista firm z filtrowaniem."""
    form = RestaurantSearchForm(request.GET or None)
    qs = Restaurant.objects.filter(is_active=True).annotate(avg_rating=Avg("reviews__rating"))

    firm_type_filter = request.GET.get("firm_type", "")

    if form.is_valid():
        firm_type = form.cleaned_data.get("firm_type")
        city = form.cleaned_data.get("city")
        min_guests = form.cleaned_data.get("min_guests")
        max_price = form.cleaned_data.get("max_price")
        has_parking = form.cleaned_data.get("has_parking")
        has_garden = form.cleaned_data.get("has_garden")
        has_dance_floor = form.cleaned_data.get("has_dance_floor")
        user_lat = form.cleaned_data.get("user_lat")
        user_lng = form.cleaned_data.get("user_lng")

        if firm_type:
            qs = qs.filter(firm_type=firm_type)
            firm_type_filter = firm_type
        if city:
            qs = qs.filter(city__icontains=city)
        event_date = form.cleaned_data.get("event_date")
        if event_date:
            booked_ids = Booking.objects.filter(
                event_date=event_date,
            ).exclude(status="cancelled").values_list("restaurant_id", flat=True)
            qs = qs.exclude(pk__in=booked_ids)
        if min_guests:
            qs = qs.filter(max_guests__gte=min_guests)
        if max_price:
            qs = qs.filter(price_per_person__lte=max_price)
        min_price = form.cleaned_data.get("min_price")
        max_guests_filter = form.cleaned_data.get("max_guests_filter")
        has_accommodation = form.cleaned_data.get("has_accommodation")
        min_rating = form.cleaned_data.get("min_rating")
        has_online_menu = form.cleaned_data.get("has_online_menu")

        if min_price:
            qs = qs.filter(price_per_person__gte=min_price)
        if max_guests_filter:
            qs = qs.filter(max_guests__lte=max_guests_filter)
        if has_parking:
            qs = qs.filter(has_parking=True)
        if has_garden:
            qs = qs.filter(has_garden=True)
        if has_dance_floor:
            qs = qs.filter(has_dance_floor=True)
        if has_accommodation:
            qs = qs.filter(has_accommodation=True)
        if min_rating:
            qs = qs.filter(avg_rating__gte=int(min_rating))
        if has_online_menu:
            qs = qs.filter(menus__is_active=True).distinct()

        attraction_type_filter = form.cleaned_data.get("attraction_type")
        if attraction_type_filter:
            qs = qs.filter(attraction_type=attraction_type_filter)

        # Filtrowanie cateringów po promieniu dowozu
        if firm_type == "catering" and user_lat and user_lng:
            nearby_ids = []
            for r in qs:
                if r.latitude and r.longitude:
                    dist = _haversine_km(
                        float(user_lat), float(user_lng),
                        float(r.latitude), float(r.longitude),
                    )
                    if dist <= r.delivery_radius_km:
                        nearby_ids.append(r.pk)
            qs = qs.filter(pk__in=nearby_ids)

    return render(request, "bookings/restaurant_list.html", {
        "restaurants": qs,
        "form": form,
        "firm_type_filter": firm_type_filter,
    })


def restaurant_detail(request, pk):
    """Szczegóły restauracji z opiniami."""
    restaurant = get_object_or_404(Restaurant, pk=pk, is_active=True)
    reviews = restaurant.reviews.select_related("user").all()
    review_form = None

    if request.user.is_authenticated:
        existing_review = Review.objects.filter(
            user=request.user, restaurant=restaurant
        ).first()
        if not existing_review:
            review_form = ReviewForm()

    # Menu widoczne publicznie (tylko z aktywnego menu)
    active_menu = Menu.objects.filter(restaurant=restaurant, is_active=True).first()
    visible_menu = MenuItem.objects.filter(
        restaurant=restaurant, is_visible=True,
        menu=active_menu,
    ) if active_menu else MenuItem.objects.none()
    categories = MenuItem.Category.choices
    menu_by_category = []
    for cat_value, cat_label in categories:
        items = visible_menu.filter(category=cat_value)
        if items.exists():
            menu_by_category.append({"label": cat_label, "items": items})

    # Czy menu jest zapisane?
    is_menu_saved = False
    if request.user.is_authenticated and menu_by_category:
        is_menu_saved = SavedMenu.objects.filter(
            user=request.user, restaurant=restaurant
        ).exists()

    # Oferta atrakcji (tylko dla typu attraction)
    attraction_items = []
    if restaurant.firm_type == Restaurant.FirmType.ATTRACTION:
        from .models import AttractionItem
        tags = AttractionItem.Tag.choices
        for tag_value, tag_label in tags:
            items = AttractionItem.objects.filter(
                restaurant=restaurant, tag=tag_value, is_active=True
            )
            if items.exists():
                attraction_items.append({"label": tag_label, "items": items})

    # Kalendarz dostępności (tylko venue z włączonym show_calendar)
    cal = None
    if restaurant.firm_type == Restaurant.FirmType.VENUE and restaurant.show_calendar:
        cal_year = request.GET.get("cal_year")
        cal_month = request.GET.get("cal_month")
        cal = build_month_grid(restaurant, cal_year, cal_month)

    return render(request, "bookings/restaurant_detail.html", {
        "restaurant": restaurant,
        "reviews": reviews,
        "review_form": review_form,
        "menu_by_category": menu_by_category,
        "is_menu_saved": is_menu_saved,
        "attraction_items": attraction_items,
        "gallery_images": restaurant.gallery_images.all(),
        "cal": cal,
    })


def restaurant_calendar_partial(request, pk):
    """AJAX endpoint – zwraca HTML kalendarza dla danego miesiąca."""
    restaurant = get_object_or_404(Restaurant, pk=pk, is_active=True)
    cal = build_month_grid(
        restaurant,
        request.GET.get("cal_year"),
        request.GET.get("cal_month"),
    )
    return render(request, "bookings/_calendar_partial.html", {"cal": cal})


# ── Rezerwacje ─────────────────────────────────────────────────────────────────

@login_required
def booking_create(request, restaurant_pk):
    """Tworzenie nowej rezerwacji."""
    restaurant = get_object_or_404(Restaurant, pk=restaurant_pk, is_active=True)
    is_attraction = restaurant.firm_type == Restaurant.FirmType.ATTRACTION

    if request.method == "POST":
        form = BookingForm(request.POST, firm_type=restaurant.firm_type)
        if form.is_valid():
            booking = form.save(commit=False)
            booking.user = request.user
            booking.restaurant = restaurant

            guest_ok = True
            if not is_attraction and booking.guest_count and booking.guest_count > restaurant.max_guests:
                form.add_error(
                    "guest_count",
                    f"Firma przyjmuje maksymalnie {restaurant.max_guests} gości.",
                )
                guest_ok = False

            if guest_ok:
                # Sprawdź czy data jest wolna
                if Booking.objects.filter(
                    restaurant=restaurant,
                    event_date=booking.event_date,
                ).exclude(status=Booking.Status.CANCELLED).exists():
                    form.add_error(
                        "event_date",
                        "Ta data jest już zarezerwowana w tej firmie.",
                    )
                else:
                    booking.save()
                    # Auto-wyślij wiadomość powitalną od właściciela firmy
                    if restaurant.welcome_message.strip():
                        owner_qs = restaurant.owners.filter(role="owner").select_related("user")
                        if not owner_qs.exists():
                            owner_qs = restaurant.owners.select_related("user")
                        if owner_qs.exists():
                            BookingMessage.objects.create(
                                booking=booking,
                                sender=owner_qs.first().user,
                                content=restaurant.welcome_message.strip(),
                            )

                    # Catering → przekieruj do ekranu wyboru menu
                    if restaurant.firm_type == Restaurant.FirmType.CATERING:
                        _active = Menu.objects.filter(restaurant=restaurant, is_active=True).first()
                        menu_items = MenuItem.objects.filter(
                            restaurant=restaurant, is_visible=True, menu=_active,
                        ) if _active else MenuItem.objects.none()
                        if menu_items.exists():
                            messages.info(
                                request,
                                "Rezerwacja złożona! Teraz wybierz pozycje z menu cateringowego.",
                            )
                            return redirect("booking_catering_menu", pk=booking.pk)

                    label = booking.get_event_type_display() if booking.event_type else "atrakcję"
                    messages.success(
                        request,
                        f"Rezerwacja na {label} została złożona! "
                        f"Oczekuj na potwierdzenie.",
                    )
                    return redirect("booking_detail", pk=booking.pk)
    else:
        initial = {
            "first_name": request.user.first_name,
            "last_name": request.user.last_name,
            "email": request.user.email,
        }
        # Auto-fill phone from user profile
        if hasattr(request.user, "profile") and request.user.profile.phone:
            initial["phone"] = request.user.profile.phone
        form = BookingForm(
            firm_type=restaurant.firm_type,
            initial=initial,
        )

    # ── Kalendarz dostępności dla formularza rezerwacji ──────────────
    cal = build_month_grid(
        restaurant,
        request.GET.get("cal_year"),
        request.GET.get("cal_month"),
        date_as_iso=True,
    )

    return render(request, "bookings/booking_form.html", {
        "form": form,
        "restaurant": restaurant,
        "is_attraction": is_attraction,
        "cal": cal,
    })


@login_required
def booking_detail(request, pk):
    """Szczegóły rezerwacji."""
    booking = get_object_or_404(Booking, pk=pk, user=request.user)

    if request.method == "POST" and request.POST.get("action") == "send_message":
        content = request.POST.get("message", "").strip()
        if content:
            BookingMessage.objects.create(
                booking=booking,
                sender=request.user,
                content=content,
            )
            messages.success(request, "Wiadomość wysłana.")
        return redirect("booking_detail", pk=booking.pk)

    chat_messages = booking.chat_messages.select_related("sender").all()
    return render(request, "bookings/booking_detail.html", {
        "booking": booking,
        "chat_messages": chat_messages,
    })


@login_required
def booking_list(request):
    """Lista rezerwacji użytkownika."""
    bookings = Booking.objects.filter(user=request.user).select_related("restaurant")
    return render(request, "bookings/booking_list.html", {"bookings": bookings})


@login_required
def booking_cancel(request, pk):
    """Anulowanie rezerwacji."""
    booking = get_object_or_404(Booking, pk=pk, user=request.user)

    if booking.status in (Booking.Status.PENDING, Booking.Status.CONFIRMED):
        booking.status = Booking.Status.CANCELLED
        booking.save()
        messages.success(request, "Rezerwacja została anulowana.")
    else:
        messages.error(request, "Nie można anulować tej rezerwacji.")

    return redirect("booking_list")


# ── Opinie ─────────────────────────────────────────────────────────────────────

@login_required
def review_create(request, restaurant_pk):
    """Dodawanie opinii o restauracji."""
    restaurant = get_object_or_404(Restaurant, pk=restaurant_pk)

    if Review.objects.filter(user=request.user, restaurant=restaurant).exists():
        messages.warning(request, "Już dodałeś opinię o tej firmie.")
        return redirect("restaurant_detail", pk=restaurant.pk)

    if request.method == "POST":
        form = ReviewForm(request.POST)
        if form.is_valid():
            review = form.save(commit=False)
            review.user = request.user
            review.restaurant = restaurant
            review.save()
            messages.success(request, "Dziękujemy za opinię!")
            return redirect("restaurant_detail", pk=restaurant.pk)
    else:
        form = ReviewForm()

    return render(request, "bookings/review_form.html", {
        "form": form,
        "restaurant": restaurant,
    })


# ── Login redirect ─────────────────────────────────────────────────────────────

@login_required
def login_redirect_view(request):
    """Przekieruj po logowaniu — firma → panel firmy, klient → strona główna."""
    if RestaurantOwner.objects.filter(user=request.user, restaurant__isnull=False).exists():
        return redirect("owner_dashboard")
    return redirect("home")


# ── Ustawienia konta ───────────────────────────────────────────────────────────

@login_required
def account_settings(request):
    """Ustawienia konta użytkownika — imię, nazwisko, email, telefon, miasto."""
    from .models import UserProfile
    profile, _ = UserProfile.objects.get_or_create(user=request.user)

    if request.method == "POST":
        form = UserSettingsForm(request.POST)
        if form.is_valid():
            request.user.first_name = form.cleaned_data["first_name"]
            request.user.last_name = form.cleaned_data["last_name"]
            request.user.email = form.cleaned_data["email"]
            request.user.save()
            profile.phone = form.cleaned_data["phone"]
            profile.city = form.cleaned_data["city"]
            profile.client_type = form.cleaned_data["client_type"]
            profile.company_name = form.cleaned_data.get("company_name", "")
            profile.company_address = form.cleaned_data.get("company_address", "")
            profile.company_nip = form.cleaned_data.get("company_nip", "")
            profile.save()
            messages.success(request, "Ustawienia konta zostały zapisane.")
            next_url = request.GET.get("next", "account_settings")
            return redirect(next_url)
    else:
        form = UserSettingsForm(initial={
            "first_name": request.user.first_name,
            "last_name": request.user.last_name,
            "email": request.user.email,
            "phone": profile.phone,
            "city": profile.city or "Poznań",
            "client_type": profile.client_type or "private",
            "company_name": profile.company_name,
            "company_address": profile.company_address,
            "company_nip": profile.company_nip,
        })

    is_new_google = request.GET.get("new") == "1"
    return render(request, "bookings/account_settings.html", {
        "form": form,
        "is_new_google": is_new_google,
    })


# ── Autoryzacja ────────────────────────────────────────────────────────────────

def register(request):
    """Rejestracja nowego użytkownika (klienta)."""
    if request.method == "POST":
        form = UserRegisterForm(request.POST)
        if form.is_valid():
            user = form.save()
            # Utwórz profil z telefonem i danymi firmowymi
            from .models import UserProfile
            UserProfile.objects.create(
                user=user,
                phone=form.cleaned_data.get("phone", ""),
                client_type=form.cleaned_data.get("client_type", "private"),
                company_name=form.cleaned_data.get("company_name", ""),
                company_address=form.cleaned_data.get("company_address", ""),
                company_nip=form.cleaned_data.get("company_nip", ""),
            )
            login(request, user, backend="django.contrib.auth.backends.ModelBackend")
            messages.success(request, f"Witaj, {user.first_name}! Konto zostało utworzone.")
            return redirect("home")
    else:
        form = UserRegisterForm()
    return render(request, "bookings/register.html", {"form": form})


def owner_register(request):
    """Rejestracja właściciela firmy."""
    if request.method == "POST":
        form = OwnerRegisterForm(request.POST)
        if form.is_valid():
            user = form.save()
            # Utwórz profil z telefonem
            from .models import UserProfile
            UserProfile.objects.create(
                user=user,
                phone=form.cleaned_data.get("phone", ""),
            )
            RestaurantOwner.objects.create(user=user, restaurant=None, role="owner")
            login(request, user, backend="django.contrib.auth.backends.ModelBackend")
            messages.success(request, f"Witaj, {user.first_name}! Konto właściciela zostało utworzone. Dodaj teraz swoją firmę.")
            return redirect("owner_restaurant_create")
    else:
        form = OwnerRegisterForm()
    return render(request, "bookings/owner_register.html", {"form": form})


def _save_gallery_images(request, restaurant):
    """Save gallery image URLs submitted from the form."""
    # Delete images marked for removal
    delete_ids = request.POST.getlist("delete_image")
    if delete_ids:
        RestaurantImage.objects.filter(
            restaurant=restaurant, pk__in=delete_ids
        ).delete()

    # Add new image URLs
    new_urls = request.POST.getlist("new_image_url")
    new_captions = request.POST.getlist("new_image_caption")
    max_order = (
        restaurant.gallery_images.aggregate(Max("order"))["order__max"] or 0
    )
    for i, url in enumerate(new_urls):
        url = url.strip()
        if url:
            caption = new_captions[i].strip() if i < len(new_captions) else ""
            max_order += 1
            RestaurantImage.objects.create(
                restaurant=restaurant,
                image_url=url,
                caption=caption,
                order=max_order,
            )

    # Add images imported from social media
    social_urls = request.POST.getlist("social_image_url")
    for url in social_urls:
        url = url.strip()
        if url and not RestaurantImage.objects.filter(restaurant=restaurant, image_url=url).exists():
            max_order += 1
            RestaurantImage.objects.create(
                restaurant=restaurant,
                image_url=url,
                caption="Import z social media",
                order=max_order,
            )


@login_required
def social_import_api(request):
    """AJAX endpoint: pobierz dane firmy z dowolnego linku."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    body = json.loads(request.body) if request.content_type == "application/json" else request.POST
    urls = body.get("urls", [])
    if isinstance(urls, str):
        urls = [urls]

    # Filter empty
    urls = [u.strip() for u in urls if u and u.strip()]

    if not urls:
        return JsonResponse({"error": "Wklej co najmniej jeden link."}, status=400)

    data = import_from_urls(urls)
    return JsonResponse(data)


@login_required
def owner_restaurant_create(request):
    """Dodawanie nowej firmy przez właściciela."""
    # Check user is an owner type at all
    if not RestaurantOwner.objects.filter(user=request.user).exists():
        messages.error(request, "Nie masz konta właściciela firmy.")
        return redirect("home")

    if request.method == "POST":
        form = RestaurantForm(request.POST)
        if form.is_valid():
            restaurant = form.save()
            # Create or update the placeholder membership
            placeholder = RestaurantOwner.objects.filter(
                user=request.user, restaurant__isnull=True
            ).first()
            if placeholder:
                placeholder.restaurant = restaurant
                placeholder.save()
            else:
                RestaurantOwner.objects.create(
                    user=request.user, restaurant=restaurant, role="owner"
                )
            _save_gallery_images(request, restaurant)
            request.session["active_restaurant_id"] = restaurant.pk
            messages.success(request, f"Firma '{restaurant.name}' została dodana!")
            return redirect("owner_dashboard")
    else:
        form = RestaurantForm()
    return render(request, "bookings/owner/restaurant_form.html", {"form": form, "editing": False})


@login_required
def owner_restaurant_edit(request):
    """Edycja aktywnej firmy."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        messages.error(request, "Nie masz konta właściciela firmy.")
        return redirect("home")
    if not restaurant:
        messages.info(request, "Najpierw dodaj swoją firmę.")
        return redirect("owner_restaurant_create")

    if request.method == "POST":
        form = RestaurantForm(request.POST, instance=restaurant)
        if form.is_valid():
            form.save()
            _save_gallery_images(request, restaurant)
            messages.success(request, "Dane firmy zostały zaktualizowane.")
            return redirect("owner_dashboard")
    else:
        form = RestaurantForm(instance=restaurant)
    return render(request, "bookings/owner/restaurant_form.html", {
        "form": form, "editing": True, "restaurant": restaurant,
        "gallery_images": restaurant.gallery_images.all(),
    })


# ── Panel właścicieli restauracji ─────────────────────────────────────────────

@login_required
def owner_dashboard(request):
    """Panel główny właściciela firmy."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        messages.error(request, "Nie masz konta właściciela firmy.")
        return redirect("home")
    if not restaurant:
        messages.info(request, "Najpierw dodaj swoją firmę.")
        return redirect("owner_restaurant_create")
    
    today = timezone.now().date()
    
    # Statystyki
    stats = {
        "total_bookings": Booking.objects.filter(restaurant=restaurant).count(),
        "pending_bookings": Booking.objects.filter(restaurant=restaurant, status=Booking.Status.PENDING).count(),
        "upcoming_bookings": Booking.objects.filter(
            restaurant=restaurant,
            event_date__gte=today,
            status=Booking.Status.CONFIRMED
        ).count(),
        "this_month_bookings": Booking.objects.filter(
            restaurant=restaurant,
            event_date__year=today.year,
            event_date__month=today.month
        ).count(),
    }
    
    # Najnowsze rezerwacje
    recent_bookings = Booking.objects.filter(restaurant=restaurant).order_by("-created_at")[:5]
    
    # Nadchodzące rezerwacje
    upcoming_bookings = Booking.objects.filter(
        restaurant=restaurant,
        event_date__gte=today,
        status=Booking.Status.CONFIRMED
    ).order_by("event_date")[:5]
    
    context = {
        "restaurant": restaurant,
        "stats": stats,
        "recent_bookings": recent_bookings,
        "upcoming_bookings": upcoming_bookings,
    }
    return render(request, "bookings/owner/dashboard.html", context)


@login_required
def owner_bookings(request):
    """Lista wszystkich rezerwacji właściciela restauracji."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        return redirect("home")
    if not restaurant:
        return redirect("owner_restaurant_create")
    
    status_filter = request.GET.get("status", "")
    bookings = Booking.objects.filter(restaurant=restaurant)
    
    if status_filter:
        bookings = bookings.filter(status=status_filter)
    
    bookings = bookings.order_by("-event_date")
    
    context = {
        "restaurant": restaurant,
        "bookings": bookings,
        "status_filter": status_filter,
        "status_choices": Booking.Status.choices,
    }
    return render(request, "bookings/owner/bookings.html", context)


@login_required
def owner_booking_detail(request, booking_id):
    """Szczegóły rezerwacji z opcją zatwierdzania/odrzucania."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        return redirect("home")
    if not restaurant:
        return redirect("owner_restaurant_create")
    
    booking = get_object_or_404(Booking, id=booking_id, restaurant=restaurant)
    
    if request.method == "POST":
        action = request.POST.get("action")
        
        if action == "confirm" and booking.status == Booking.Status.PENDING:
            booking.status = Booking.Status.CONFIRMED
            booking.save()
            messages.success(request, f"Rezerwacja #{booking.id} została potwierdzona.")
            
        elif action == "cancel" and booking.status in (Booking.Status.PENDING, Booking.Status.CONFIRMED):
            booking.status = Booking.Status.CANCELLED
            booking.save()
            messages.success(request, f"Rezerwacja #{booking.id} została anulowana.")

        elif action == "add_note":
            note_date = request.POST.get("note_date")
            note_title = request.POST.get("note_title", "").strip()
            note_content = request.POST.get("note_content", "").strip()
            if note_date and note_title and note_content:
                BookingNote.objects.create(
                    booking=booking,
                    author=request.user,
                    date=note_date,
                    title=note_title,
                    content=note_content,
                )
                messages.success(request, "Notatka została dodana.")
            else:
                messages.error(request, "Wypełnij wszystkie pola notatki.")

        elif action == "delete_note":
            note_id = request.POST.get("note_id")
            BookingNote.objects.filter(id=note_id, booking=booking).delete()
            messages.success(request, "Notatka została usunięta.")

        elif action == "add_todo":
            todo_text = request.POST.get("todo_text", "").strip()
            if todo_text:
                BookingTodo.objects.create(booking=booking, text=todo_text)
                messages.success(request, "Zadanie dodane.")
            else:
                messages.error(request, "Wpisz treść zadania.")

        elif action == "toggle_todo":
            todo_id = request.POST.get("todo_id")
            todo = BookingTodo.objects.filter(id=todo_id, booking=booking).first()
            if todo:
                todo.is_done = not todo.is_done
                todo.save()

        elif action == "delete_todo":
            todo_id = request.POST.get("todo_id")
            BookingTodo.objects.filter(id=todo_id, booking=booking).delete()
            messages.success(request, "Zadanie usunięte.")

        elif action == "send_message":
            content = request.POST.get("message", "").strip()
            if content:
                BookingMessage.objects.create(
                    booking=booking,
                    sender=request.user,
                    content=content,
                )
                messages.success(request, "Wiadomość wysłana.")

        elif action == "close_deal":
            agreed_price = request.POST.get("deal_agreed_price", "").strip()
            deal_terms = request.POST.get("deal_terms", "").strip()
            if agreed_price:
                try:
                    from decimal import Decimal
                    booking.deal_agreed_price = Decimal(agreed_price)
                    booking.deal_terms = deal_terms
                    booking.deal_closed_at = timezone.now()
                    booking.status = Booking.Status.CONFIRMED
                    booking.save()
                    messages.success(request, "Deal zamknięty! Umowa została wygenerowana.")
                    return redirect("owner_booking_agreement", booking_id=booking.id)
                except Exception:
                    messages.error(request, "Nieprawidłowa cena.")
            else:
                messages.error(request, "Podaj uzgodnioną cenę.")

        elif action == "reopen_deal":
            booking.deal_closed_at = None
            booking.deal_agreed_price = None
            booking.deal_terms = ""
            booking.save()
            messages.success(request, "Deal został otwarty ponownie.")
        
        return redirect("owner_booking_detail", booking_id=booking.id)

    notes = booking.crm_notes.all()
    chat_messages = booking.chat_messages.select_related("sender").all()
    todos = booking.todos.all()
    todos_done = todos.filter(is_done=True).count()
    today = timezone.now().date().isoformat()
    
    return render(request, "bookings/owner/booking_detail.html", {
        "restaurant": restaurant,
        "booking": booking,
        "notes": notes,
        "chat_messages": chat_messages,
        "todos": todos,
        "todos_done": todos_done,
        "today": today,
    })


@login_required
def owner_booking_agreement(request, booking_id):
    """Wyświetlenie umowy / potwierdzenia dealu — do druku / PDF."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        return redirect("home")
    if not restaurant:
        return redirect("owner_restaurant_create")

    booking = get_object_or_404(Booking, id=booking_id, restaurant=restaurant)

    if not booking.deal_closed_at:
        messages.error(request, "Deal nie został jeszcze zamknięty.")
        return redirect("owner_booking_detail", booking_id=booking.id)

    menu_selections = booking.menu_selections.select_related("menu_item").all()
    menu_total = sum(s.subtotal for s in menu_selections)

    return render(request, "bookings/owner/agreement.html", {
        "restaurant": restaurant,
        "booking": booking,
        "menu_selections": menu_selections,
        "menu_total": menu_total,
    })


@login_required
def owner_calendar(request):
    """Kalendarz rezerwacji restauracji — widok miesięczny (AJAX-ready)."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        messages.error(request, "Nie masz uprawnień do zarządzania firmą.")
        return redirect("home")
    if not restaurant:
        return redirect("owner_restaurant_create")

    today = timezone.now().date()
    year = request.GET.get("year", today.year)
    month = request.GET.get("month", today.month)

    cal = build_month_grid(restaurant, year, month, include_bookings=True)

    ctx = {
        "restaurant": restaurant,
        "cal": cal,
    }

    # AJAX request → return only the calendar partial
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return render(request, "bookings/owner/_calendar_grid.html", ctx)

    return render(request, "bookings/owner/calendar.html", ctx)


@login_required
def owner_calendar_api(request):
    """JSON API: calendar data, day details, block/unblock dates."""
    memberships, restaurant, membership = _get_owner_context(request)
    if not restaurant:
        return JsonResponse({"error": "Brak firmy"}, status=404)

    action = request.GET.get("action") or (
        json.loads(request.body).get("action") if request.method == "POST" and request.content_type == "application/json" else request.POST.get("action")
    )

    # ── GET: month grid data ─────────────────────────────────────────────
    if request.method == "GET" and action in (None, "month"):
        year = request.GET.get("year")
        month = request.GET.get("month")
        cal = build_month_grid(restaurant, year, month, include_bookings=True, date_as_iso=True)
        return JsonResponse(cal, safe=False)

    # ── GET: day detail ──────────────────────────────────────────────────
    if request.method == "GET" and action == "day":
        try:
            d = date.fromisoformat(request.GET["date"])
        except (KeyError, ValueError):
            return JsonResponse({"error": "Podaj datę w formacie YYYY-MM-DD"}, status=400)
        detail = get_day_details(restaurant, d)
        detail["date"] = detail["date"].isoformat()
        detail["blocked"] = {
            "is_blocked": detail["blocked"] is not None,
            "reason": detail["blocked"].reason if detail["blocked"] else "",
        }
        for b in detail["bookings"]:
            b["start_time"] = b["start_time"].strftime("%H:%M") if b["start_time"] else None
            b["end_time"] = b["end_time"].strftime("%H:%M") if b["end_time"] else None
            b["created_at"] = b["created_at"].isoformat() if b["created_at"] else None
        return JsonResponse(detail)

    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    body = json.loads(request.body) if request.content_type == "application/json" else request.POST

    # ── POST: block date ─────────────────────────────────────────────────
    if action == "block":
        try:
            d = date.fromisoformat(body.get("date", ""))
        except ValueError:
            return JsonResponse({"error": "Nieprawidłowa data"}, status=400)
        reason = body.get("reason", "").strip()
        obj, created = BlockedDate.objects.get_or_create(
            restaurant=restaurant, date=d,
            defaults={"reason": reason},
        )
        if not created:
            obj.reason = reason
            obj.save()
        return JsonResponse({"ok": True, "created": created, "date": d.isoformat()})

    # ── POST: unblock date ───────────────────────────────────────────────
    if action == "unblock":
        try:
            d = date.fromisoformat(body.get("date", ""))
        except ValueError:
            return JsonResponse({"error": "Nieprawidłowa data"}, status=400)
        deleted, _ = BlockedDate.objects.filter(restaurant=restaurant, date=d).delete()
        return JsonResponse({"ok": True, "deleted": bool(deleted), "date": d.isoformat()})

    return JsonResponse({"error": "Nieznana akcja"}, status=400)


# ── Menu restauracji (panel właściciela) ──────────────────────────────────────────

@login_required
def owner_menu(request):
    """Lista menu restauracji."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        return redirect("home")
    if not restaurant:
        return redirect("owner_restaurant_create")

    all_menus = Menu.objects.filter(restaurant=restaurant).annotate(
        item_count=Count("items"),
    ).select_related("last_edited_by")

    if request.method == "POST":
        action = request.POST.get("action")

        # ── Zapis włączonych typów menu ──
        if action == "save_menu_types":
            enabled = request.POST.getlist("menu_types")
            valid_keys = {t["key"] for t in MENU_TYPE_INFO}
            restaurant.enabled_menu_types = [k for k in enabled if k in valid_keys]
            restaurant.save(update_fields=["enabled_menu_types"])
            messages.success(request, "Dost\u0119pne typy menu zosta\u0142y zaktualizowane.")
            return redirect("owner_menu")

        # ── Nowe menu ──
        if action == "create_menu":
            menu_name = request.POST.get("menu_name", "").strip()
            menu_desc = request.POST.get("menu_description", "").strip()
            if menu_name:
                new_menu = Menu.objects.create(
                    restaurant=restaurant,
                    name=menu_name,
                    description=menu_desc,
                    last_edited_by=request.user,
                )
                messages.success(request, f'Utworzono menu „{menu_name}".')
                return redirect("owner_menu_detail", menu_id=new_menu.id)
            else:
                messages.error(request, "Podaj nazwę menu.")

        # ── Ustaw jako aktualne ──
        elif action == "set_active":
            menu_id = request.POST.get("menu_id")
            menu_obj = Menu.objects.filter(id=menu_id, restaurant=restaurant).first()
            if menu_obj:
                Menu.objects.filter(restaurant=restaurant).update(is_active=False)
                menu_obj.is_active = True
                menu_obj.save(update_fields=["is_active"])
                messages.success(request, f'Menu „{menu_obj.name}" jest teraz aktualnym menu.')

        # ── Kopiuj menu ──
        elif action == "copy_menu":
            src_id = request.POST.get("menu_id")
            src = Menu.objects.filter(id=src_id, restaurant=restaurant).first()
            if src:
                new_menu = Menu.objects.create(
                    restaurant=restaurant,
                    name=f"Kopia {src.name}",
                    description=src.description,
                    last_edited_by=request.user,
                )
                # Skopiuj pozycje
                for item in MenuItem.objects.filter(menu=src):
                    MenuItem.objects.create(
                        restaurant=restaurant,
                        menu=new_menu,
                        category=item.category,
                        name=item.name,
                        description=item.description,
                        price=item.price,
                        is_visible=item.is_visible,
                        order=item.order,
                    )
                messages.success(request, f'Utworzono kopię menu „{src.name}" ({new_menu.items.count()} pozycji).')
                return redirect("owner_menu_detail", menu_id=new_menu.id)

        # ── Usuń menu ──
        elif action == "delete_menu":
            menu_id = request.POST.get("menu_id")
            menu_obj = Menu.objects.filter(id=menu_id, restaurant=restaurant).first()
            if menu_obj:
                if menu_obj.is_active:
                    messages.error(request, "Nie można usunąć aktualnego menu. Najpierw ustaw inne jako aktualne.")
                else:
                    name = menu_obj.name
                    menu_obj.delete()
                    messages.success(request, f'Usunięto menu „{name}".')

        return redirect("owner_menu")

    # Jeśli brak menu — utwórz domyślne
    if not all_menus.exists():
        Menu.objects.create(
            restaurant=restaurant,
            name="Menu główne",
            is_active=True,
            last_edited_by=request.user,
        )
        all_menus = Menu.objects.filter(restaurant=restaurant).annotate(
            item_count=Count("items"),
        ).select_related("last_edited_by")

    return render(request, "bookings/owner/menu_list.html", {
        "restaurant": restaurant,
        "all_menus": all_menus,
        "menu_type_info": MENU_TYPE_INFO,
        "enabled_menu_types": restaurant.enabled_menu_types or [],
    })


@login_required
def owner_menu_detail(request, menu_id):
    """Szczegóły i edycja konkretnego menu."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        return redirect("home")
    if not restaurant:
        return redirect("owner_restaurant_create")

    current_menu = get_object_or_404(Menu, id=menu_id, restaurant=restaurant)

    def _upsert_template(category, name, price):
        try:
            tpl, created = MenuItemTemplate.objects.get_or_create(
                category=category, name=name,
                defaults={"last_price": price},
            )
            if not created:
                tpl.last_price = price
                tpl.usage_count = tpl.usage_count + 1
                tpl.save(update_fields=["last_price", "usage_count", "updated_at"])
        except Exception:
            pass

    def _touch():
        current_menu.last_edited_by = request.user
        current_menu.save(update_fields=["last_edited_by", "last_edited_at"])

    if request.method == "POST":
        action = request.POST.get("action")

        # ── Ustaw jako aktualne ──
        if action == "set_active":
            Menu.objects.filter(restaurant=restaurant).update(is_active=False)
            current_menu.is_active = True
            current_menu.save(update_fields=["is_active"])
            messages.success(request, f'Menu „{current_menu.name}" jest teraz aktualnym menu.')

        # ── Zmień nazwę / opis ──
        elif action == "edit_meta":
            new_name = request.POST.get("menu_name", "").strip()
            new_desc = request.POST.get("menu_description", "").strip()
            if new_name:
                current_menu.name = new_name
                current_menu.description = new_desc
                current_menu.save(update_fields=["name", "description"])
                _touch()
                messages.success(request, "Zaktualizowano dane menu.")

        # ── Usuń menu ──
        elif action == "delete_menu":
            if current_menu.is_active:
                messages.error(request, "Nie można usunąć aktualnego menu. Najpierw ustaw inne jako aktualne.")
            else:
                name = current_menu.name
                current_menu.delete()
                messages.success(request, f'Usunięto menu „{name}".')
                return redirect("owner_menu")

        # ── Dodaj pozycję ──
        elif action == "add":
            category = request.POST.get("category", "")
            name = request.POST.get("name", "").strip()
            description = request.POST.get("description", "").strip()
            price = request.POST.get("price", "0")
            is_visible = request.POST.get("is_visible") == "on"
            if name and category:
                try:
                    MenuItem.objects.create(
                        restaurant=restaurant,
                        menu=current_menu,
                        category=category,
                        name=name,
                        description=description,
                        price=price,
                        is_visible=is_visible,
                    )
                    _upsert_template(category, name, price)
                    _touch()
                    messages.success(request, f"Dodano: {name}")
                except Exception:
                    messages.error(request, "Błąd przy dodawaniu pozycji.")
            else:
                messages.error(request, "Podaj nazwę i kategorię.")

        # ── Edytuj pozycję ──
        elif action == "edit":
            item_id = request.POST.get("item_id")
            item = MenuItem.objects.filter(id=item_id, menu=current_menu).first()
            if item:
                item.category = request.POST.get("category", item.category)
                item.name = request.POST.get("name", item.name).strip()
                item.description = request.POST.get("description", "").strip()
                item.price = request.POST.get("price", item.price)
                item.is_visible = request.POST.get("is_visible") == "on"
                item.save()
                _upsert_template(item.category, item.name, item.price)
                _touch()
                messages.success(request, f"Zaktualizowano: {item.name}")

        # ── Usuń pozycję ──
        elif action == "delete":
            item_id = request.POST.get("item_id")
            MenuItem.objects.filter(id=item_id, menu=current_menu).delete()
            _touch()
            messages.success(request, "Pozycja usunięta.")

        return redirect("owner_menu_detail", menu_id=current_menu.id)

    # Grupuj pozycje wg kategorii
    categories = MenuItem.Category.choices
    menu_by_category = []
    for cat_value, cat_label in categories:
        items = MenuItem.objects.filter(menu=current_menu, category=cat_value)
        menu_by_category.append({
            "value": cat_value,
            "label": cat_label,
            "items": items,
            "count": items.count(),
        })

    return render(request, "bookings/owner/menu_detail.html", {
        "restaurant": restaurant,
        "menu_by_category": menu_by_category,
        "categories": categories,
        "current_menu": current_menu,
    })


# ── Wybór menu przez klienta ────────────────────────────────────────────────

def _build_menu_by_category(menu_items, booking):
    """Przygotuj dane menu pogrupowane wg kategorii z aktualnymi wyborami."""
    current = {}
    for s in booking.menu_selections.all():
        current[s.menu_item_id] = {"qty": s.quantity, "serving_order": s.serving_order, "guest_group": s.guest_group}
    categories = MenuItem.Category.choices
    result = []
    for cat_value, cat_label in categories:
        items = menu_items.filter(category=cat_value)
        items_data = []
        for item in items:
            sel = current.get(item.id, {})
            items_data.append({
                "item": item,
                "qty": sel.get("qty", 0),
                "selected": item.id in current,
                "serving_order": sel.get("serving_order", 0),
                "guest_group": sel.get("guest_group", ""),
            })
        if items_data:
            result.append({"label": cat_label, "value": cat_value, "items": items_data})
    return result


def _log_menu_crm(booking, user, new_selections, old_selections, menu_type_label=""):
    """Loguj zmiany menu w CRM."""
    prefix = f"[{menu_type_label}] " if menu_type_label else ""
    if new_selections:
        changes_text = "\n".join(new_selections)
        note_title = f"{prefix}Zmiana wyboru menu" if old_selections else f"{prefix}Wybór menu"
        BookingNote.objects.create(
            booking=booking,
            author=user,
            date=timezone.now().date(),
            title=note_title,
            content=f"Klient zaktualizował wybór menu:\n{changes_text}",
        )
    elif old_selections:
        BookingNote.objects.create(
            booking=booking,
            author=user,
            date=timezone.now().date(),
            title=f"{prefix}Usunięcie wyboru menu",
            content="Klient usunął wszystkie pozycje z menu.",
        )


@login_required
def booking_menu_select(request, pk):
    """Wybór/edycja menu przez klienta po rezerwacji — router po typie menu."""
    booking = get_object_or_404(Booking, pk=pk, user=request.user)
    restaurant = booking.restaurant
    active_menu = Menu.objects.filter(restaurant=restaurant, is_active=True).first()
    menu_items = MenuItem.objects.filter(
        restaurant=restaurant, is_visible=True, menu=active_menu,
    ) if active_menu else MenuItem.objects.none()

    if not menu_items.exists():
        messages.info(request, "Ta firma nie udostępniła jeszcze menu.")
        return redirect("booking_detail", pk=pk)

    # ── Określ dostępne typy menu ──
    if restaurant.firm_type == Restaurant.FirmType.VENUE:
        enabled_types = restaurant.enabled_menu_types or ["detailed"]
    else:
        enabled_types = ["detailed"]

    # ── Wybór typu menu (jeśli potrzebny) ──
    if not booking.menu_type:
        if len(enabled_types) > 1:
            # Pokaż selektor typów
            if request.method == "POST" and request.POST.get("action") == "select_type":
                selected = request.POST.get("menu_type")
                if selected in enabled_types:
                    booking.menu_type = selected
                    booking.save(update_fields=["menu_type"])
                    return redirect("booking_menu_select", pk=pk)
            type_info = [t for t in MENU_TYPE_INFO if t["key"] in enabled_types]
            return render(request, "bookings/booking_menu_type.html", {
                "booking": booking,
                "restaurant": restaurant,
                "menu_type_info": type_info,
            })
        else:
            booking.menu_type = enabled_types[0]
            booking.save(update_fields=["menu_type"])

    menu_type = booking.menu_type or "detailed"
    menu_by_category = _build_menu_by_category(menu_items, booking)
    type_label = MENU_TYPE_MAP.get(menu_type, {}).get("label", "")
    can_change_type = len(enabled_types) > 1

    # ── Zmiana typu menu ──
    if request.method == "POST" and request.POST.get("action") == "change_type":
        booking.menu_type = ""
        booking.menu_selections.all().delete()
        booking.courses.all().delete()
        booking.save(update_fields=["menu_type"])
        return redirect("booking_menu_select", pk=pk)

    # ── POST: zapis wg typu ──
    if request.method == "POST" and request.POST.get("action") != "change_type":
        old_selections = {s.menu_item_id: s.quantity for s in booking.menu_selections.all()}

        if menu_type == "detailed":
            return _save_menu_detailed(request, booking, menu_items, old_selections, type_label)
        elif menu_type == "limited":
            return _save_menu_limited(request, booking, menu_items, old_selections, type_label)
        elif menu_type == "buffet":
            return _save_menu_buffet(request, booking, menu_items, old_selections, type_label)
        elif menu_type == "custom_mix":
            return _save_menu_custom_mix(request, booking, menu_items, old_selections, type_label)

    # ── GET: pokaz template wg typu ──
    base_ctx = {
        "booking": booking,
        "restaurant": restaurant,
        "menu_by_category": menu_by_category,
        "menu_type": menu_type,
        "type_label": type_label,
        "type_info": MENU_TYPE_MAP.get(menu_type, {}),
        "can_change_type": can_change_type,
    }

    if menu_type == "limited":
        return render(request, "bookings/booking_menu_limited.html", base_ctx)
    elif menu_type == "buffet":
        return render(request, "bookings/booking_menu_buffet.html", base_ctx)
    elif menu_type == "custom_mix":
        base_ctx["courses"] = booking.courses.prefetch_related("items__menu_item").all()
        base_ctx["serving_styles"] = BookingCourse.ServingStyle.choices
        return render(request, "bookings/booking_menu_custom.html", base_ctx)
    else:
        return render(request, "bookings/booking_menu.html", base_ctx)


def _save_menu_detailed(request, booking, menu_items, old_selections, type_label):
    """Zapis menu: Szczegółowy plan (ilości)."""
    booking.menu_selections.all().delete()
    new_selections = []
    for item in menu_items:
        qty_str = request.POST.get(f"qty_{item.id}", "0")
        try:
            qty = int(qty_str)
        except ValueError:
            qty = 0
        if qty > 0:
            BookingMenuItem.objects.create(booking=booking, menu_item=item, quantity=qty)
            new_selections.append(f"{item.name} x{qty}")

    _log_menu_crm(booking, request.user, new_selections, old_selections, type_label)
    if new_selections:
        messages.success(request, "Menu zostało zapisane.")
    else:
        messages.info(request, "Nie wybrano żadnych pozycji menu.")
    return redirect("booking_detail", pk=booking.pk)


def _save_menu_limited(request, booking, menu_items, old_selections, type_label):
    """Zapis menu: Ograniczony Wybór (zaznaczone pozycje, bez ilości)."""
    booking.menu_selections.all().delete()
    new_selections = []
    for item in menu_items:
        if request.POST.get(f"sel_{item.id}"):
            BookingMenuItem.objects.create(booking=booking, menu_item=item, quantity=1)
            new_selections.append(f"{item.name} (na karcie)")

    _log_menu_crm(booking, request.user, new_selections, old_selections, type_label)
    if new_selections:
        messages.success(request, "Karta menu została zapisana.")
    else:
        messages.info(request, "Nie wybrano żadnych pozycji do karty.")
    return redirect("booking_detail", pk=booking.pk)


def _save_menu_buffet(request, booking, menu_items, old_selections, type_label):
    """Zapis menu: Szwedzki Stół (zaznaczone pozycje + kolejność)."""
    booking.menu_selections.all().delete()
    new_selections = []
    order_counter = 1
    # Items come ordered by their serving_order inputs
    ordered_ids = request.POST.get("buffet_order", "").split(",")
    for item_id_str in ordered_ids:
        try:
            item_id = int(item_id_str.strip())
        except (ValueError, AttributeError):
            continue
        item = menu_items.filter(id=item_id).first()
        if item and request.POST.get(f"sel_{item.id}"):
            BookingMenuItem.objects.create(
                booking=booking, menu_item=item, quantity=1, serving_order=order_counter,
            )
            new_selections.append(f"{order_counter}. {item.name}")
            order_counter += 1
    # Also handle checked items not in the order list
    for item in menu_items:
        if request.POST.get(f"sel_{item.id}") and not booking.menu_selections.filter(menu_item=item).exists():
            BookingMenuItem.objects.create(
                booking=booking, menu_item=item, quantity=1, serving_order=order_counter,
            )
            new_selections.append(f"{order_counter}. {item.name}")
            order_counter += 1

    _log_menu_crm(booking, request.user, new_selections, old_selections, type_label)
    if new_selections:
        messages.success(request, "Menu bufetowe zostało zapisane.")
    else:
        messages.info(request, "Nie wybrano żadnych pozycji na stół.")
    return redirect("booking_detail", pk=booking.pk)


def _save_menu_custom_mix(request, booking, menu_items, old_selections, type_label):
    """Zapis menu: Własny Mix (etapy + pozycje z grupami gości)."""
    action = request.POST.get("action", "")

    # Dodaj nowy etap
    if action == "add_course":
        course_name = request.POST.get("course_name", "").strip()
        serving_style = request.POST.get("serving_style", "for_all")
        if course_name:
            max_order = booking.courses.aggregate(m=Max("order"))["m"] or 0
            BookingCourse.objects.create(
                booking=booking, name=course_name, order=max_order + 1,
                serving_style=serving_style,
            )
            messages.success(request, f'Dodano etap „{course_name}".')
        return redirect("booking_menu_select", pk=booking.pk)

    # Usuń etap
    if action == "delete_course":
        course_id = request.POST.get("course_id")
        booking.courses.filter(id=course_id).delete()
        messages.success(request, "Etap usunięty.")
        return redirect("booking_menu_select", pk=booking.pk)

    # Dodaj pozycję do etapu
    if action == "add_course_item":
        course_id = request.POST.get("course_id")
        item_id = request.POST.get("item_id")
        guest_group = request.POST.get("guest_group", "").strip()
        course = booking.courses.filter(id=course_id).first()
        item = menu_items.filter(id=item_id).first()
        if course and item:
            # Usuń z innych etapów jeśli istnieje
            booking.menu_selections.filter(menu_item=item).delete()
            max_order = course.items.aggregate(m=Max("serving_order"))["m"] or 0
            BookingMenuItem.objects.create(
                booking=booking, menu_item=item, quantity=1,
                course=course, guest_group=guest_group,
                serving_order=max_order + 1,
            )
        return redirect("booking_menu_select", pk=booking.pk)

    # Usuń pozycję z etapu
    if action == "remove_course_item":
        sel_id = request.POST.get("selection_id")
        booking.menu_selections.filter(id=sel_id).delete()
        return redirect("booking_menu_select", pk=booking.pk)

    # Finalizuj cały plan
    if action == "finalize_mix":
        new_selections = []
        for course in booking.courses.all():
            new_selections.append(f"— {course.name} ({course.get_serving_style_display()}) —")
            for sel in course.items.select_related("menu_item").all():
                grp = f" [{sel.guest_group}]" if sel.guest_group else ""
                new_selections.append(f"  {sel.menu_item.name}{grp}")
        _log_menu_crm(booking, request.user, new_selections, old_selections, type_label)
        messages.success(request, "Plan menu został zapisany.")
        return redirect("booking_detail", pk=booking.pk)

    return redirect("booking_menu_select", pk=booking.pk)


@login_required
def booking_catering_menu(request, pk):
    """Ekran wyboru menu cateringowego — wyświetlany po złożeniu rezerwacji na catering."""
    booking = get_object_or_404(Booking, pk=pk, user=request.user)
    restaurant = booking.restaurant

    if restaurant.firm_type != Restaurant.FirmType.CATERING:
        return redirect("booking_detail", pk=pk)

    active_menu = Menu.objects.filter(restaurant=restaurant, is_active=True).first()
    menu_items = MenuItem.objects.filter(
        restaurant=restaurant, is_visible=True, menu=active_menu,
    ) if active_menu else MenuItem.objects.none()

    if request.method == "POST":
        # Zapisz wybory menu
        booking.menu_selections.all().delete()

        new_selections = []
        for item in menu_items:
            qty_str = request.POST.get(f"qty_{item.id}", "0")
            try:
                qty = int(qty_str)
            except ValueError:
                qty = 0
            if qty > 0:
                BookingMenuItem.objects.create(
                    booking=booking,
                    menu_item=item,
                    quantity=qty,
                )
                new_selections.append(f"{item.name} x{qty}")

        if new_selections:
            changes_text = "\n".join(new_selections)
            BookingNote.objects.create(
                booking=booking,
                author=request.user,
                date=timezone.now().date(),
                title="Wybór menu cateringowego",
                content=f"Klient wybrał pozycje z menu:\n{changes_text}",
            )
            messages.success(request, "Menu cateringowe zostało zapisane! Rezerwacja złożona.")
        else:
            messages.success(request, "Rezerwacja złożona bez wyboru menu.")

        return redirect("booking_detail", pk=pk)

    # Przygotuj dane menu pogrupowane po kategoriach
    current_selections = {s.menu_item_id: s.quantity for s in booking.menu_selections.all()}
    categories = MenuItem.Category.choices
    menu_by_category = []
    for cat_value, cat_label in categories:
        items = menu_items.filter(category=cat_value)
        items_with_qty = []
        for item in items:
            items_with_qty.append({
                "item": item,
                "qty": current_selections.get(item.id, 0),
            })
        if items_with_qty:
            menu_by_category.append({
                "label": cat_label,
                "items": items_with_qty,
            })

    return render(request, "bookings/booking_catering_menu.html", {
        "booking": booking,
        "restaurant": restaurant,
        "menu_by_category": menu_by_category,
    })


# ── Oferta atrakcji ─────────────────────────────────────────────────────────

@login_required
def owner_attractions(request):
    """Zarządzanie ofertą atrakcji."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        return redirect("home")
    if not restaurant:
        return redirect("owner_restaurant_create")

    if restaurant.firm_type != Restaurant.FirmType.ATTRACTION:
        messages.error(request, "Ta zakładka jest dostępna tylko dla firm typu Atrakcje.")
        return redirect("owner_dashboard")

    if request.method == "POST":
        action = request.POST.get("action")

        if action == "add":
            name = request.POST.get("name", "").strip()
            description = request.POST.get("description", "").strip()
            image_url = request.POST.get("image_url", "").strip()
            price = request.POST.get("price", "0")
            tag = request.POST.get("tag", "")
            if name and tag:
                try:
                    AttractionItem.objects.create(
                        restaurant=restaurant,
                        name=name,
                        description=description,
                        image_url=image_url,
                        price=price,
                        tag=tag,
                    )
                    messages.success(request, f"Dodano: {name}")
                except Exception:
                    messages.error(request, "Błąd przy dodawaniu pozycji.")
            else:
                messages.error(request, "Podaj nazwę i tag.")

        elif action == "edit":
            item_id = request.POST.get("item_id")
            item = AttractionItem.objects.filter(id=item_id, restaurant=restaurant).first()
            if item:
                item.name = request.POST.get("name", item.name).strip()
                item.description = request.POST.get("description", "").strip()
                item.image_url = request.POST.get("image_url", "").strip()
                item.price = request.POST.get("price", item.price)
                item.tag = request.POST.get("tag", item.tag)
                item.is_active = request.POST.get("is_active") == "on"
                item.save()
                messages.success(request, f"Zaktualizowano: {item.name}")

        elif action == "delete":
            item_id = request.POST.get("item_id")
            AttractionItem.objects.filter(id=item_id, restaurant=restaurant).delete()
            messages.success(request, "Pozycja usunięta.")

        return redirect("owner_attractions")

    tags = AttractionItem.Tag.choices
    items_by_tag = []
    for tag_value, tag_label in tags:
        items = AttractionItem.objects.filter(restaurant=restaurant, tag=tag_value)
        items_by_tag.append({
            "value": tag_value,
            "label": tag_label,
            "items": items,
            "count": items.count(),
        })

    return render(request, "bookings/owner/attractions.html", {
        "restaurant": restaurant,
        "items_by_tag": items_by_tag,
        "tags": tags,
    })


# ── Zarządzanie firmami (multi-firma) ──────────────────────────────────────────

@login_required
def owner_firms(request):
    """Lista firm użytkownika, przełączanie aktywnej firmy, zarządzanie pracownikami."""
    if not RestaurantOwner.objects.filter(user=request.user).exists():
        messages.error(request, "Nie masz konta właściciela firmy.")
        return redirect("home")

    memberships = RestaurantOwner.objects.filter(
        user=request.user, restaurant__isnull=False,
    ).select_related("restaurant")

    active_id = request.session.get("active_restaurant_id")

    # Handle POST actions
    if request.method == "POST":
        action = request.POST.get("action")

        if action == "switch":
            new_id = request.POST.get("restaurant_id")
            if memberships.filter(restaurant_id=new_id).exists():
                request.session["active_restaurant_id"] = int(new_id)
                messages.success(request, "Przełączono aktywną firmę.")
            return redirect("owner_firms")

        elif action == "remove_worker":
            worker_id = request.POST.get("membership_id")
            # Only owners can remove workers
            membership_to_delete = RestaurantOwner.objects.filter(
                id=worker_id, role="worker",
            ).first()
            if membership_to_delete:
                # Verify current user is owner of that restaurant
                is_owner = RestaurantOwner.objects.filter(
                    user=request.user,
                    restaurant=membership_to_delete.restaurant,
                    role="owner",
                ).exists()
                if is_owner:
                    membership_to_delete.delete()
                    messages.success(request, "Pracownik usunięty z firmy.")
            return redirect("owner_firms")

        elif action == "add_worker":
            restaurant_id = request.POST.get("restaurant_id")
            worker_username = request.POST.get("worker_username", "").strip()
            # Verify user is owner of this restaurant
            is_owner = RestaurantOwner.objects.filter(
                user=request.user, restaurant_id=restaurant_id, role="owner",
            ).exists()
            if is_owner and worker_username:
                from django.contrib.auth.models import User as UserModel
                try:
                    worker_user = UserModel.objects.get(username=worker_username)
                    _, created = RestaurantOwner.objects.get_or_create(
                        user=worker_user,
                        restaurant_id=restaurant_id,
                        defaults={"role": "worker"},
                    )
                    if created:
                        messages.success(request, f"Dodano pracownika: {worker_user.get_full_name() or worker_username}")
                    else:
                        messages.info(request, "Ten użytkownik jest już członkiem firmy.")
                except UserModel.DoesNotExist:
                    messages.error(request, f"Nie znaleziono użytkownika: {worker_username}")
            return redirect("owner_firms")

    # Build context: list of firms with their workers
    firms_data = []
    for m in memberships:
        all_members = RestaurantOwner.objects.filter(
            restaurant=m.restaurant,
        ).select_related("user")
        firms_data.append({
            "membership": m,
            "restaurant": m.restaurant,
            "is_active": m.restaurant_id == active_id,
            "is_owner": m.role == "owner",
            "members": all_members,
        })

    return render(request, "bookings/owner/firms.html", {
        "firms_data": firms_data,
        "active_id": active_id,
    })


@login_required
def owner_switch_firm(request, restaurant_id):
    """Quick switch active firm (GET-based for navbar dropdown)."""
    membership = RestaurantOwner.objects.filter(
        user=request.user, restaurant_id=restaurant_id,
    ).first()
    if membership:
        request.session["active_restaurant_id"] = restaurant_id
    return redirect(request.GET.get("next", "owner_dashboard"))


# ── Moja Baza Potraw ──────────────────────────────────────────────────────────

DISH_SOURCE_LABELS = dict(Dish.Source.choices)

DISH_SOURCE_ICONS = {
    "manual": "bi-pencil",
    "excel": "bi-file-earmark-excel",
    "word": "bi-file-earmark-word",
    "text": "bi-textarea-t",
    "external": "bi-globe2",
    "menu": "bi-journal-text",
}


def _parse_dishes_from_text(raw_text):
    """Parse raw text lines into dish dicts.

    Supported formats per line:
        Nazwa potrawy - opis - 45.00
        Nazwa potrawy  45.00
        Nazwa potrawy - 45
        Nazwa potrawy
    Returns list of dicts: {name, description, price}
    """
    import re
    dishes = []
    for line in raw_text.strip().splitlines():
        line = line.strip()
        if not line or len(line) < 2:
            continue
        # Try to extract price at the end
        price_match = re.search(r'[\s\-–—]+(\d+(?:[.,]\d{1,2})?)\s*(zł|PLN|pln|zl)?\s*$', line)
        price = 0
        if price_match:
            price = float(price_match.group(1).replace(',', '.'))
            line = line[:price_match.start()].strip()
        # Try to split name - description by " - " or " – "
        parts = re.split(r'\s+[-–—]\s+', line, maxsplit=1)
        name = parts[0].strip().rstrip('-–— ')
        description = parts[1].strip() if len(parts) > 1 else ''
        if name:
            dishes.append({'name': name, 'description': description, 'price': price})
    return dishes


def _parse_dishes_from_excel(file_obj):
    """Parse uploaded Excel (.xlsx) file into dish dicts."""
    import openpyxl
    wb = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
    dishes = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(min_row=1, values_only=True):
            cells = [c for c in row if c is not None]
            if not cells:
                continue
            name = str(cells[0]).strip()
            if not name or name.lower() in ('nazwa', 'name', 'potrawa', 'danie', 'lp', 'lp.'):
                continue
            description = ''
            price = 0
            for c in cells[1:]:
                val = str(c).strip()
                try:
                    price = float(val.replace(',', '.').replace(' ', ''))
                except (ValueError, AttributeError):
                    if not description:
                        description = val
            dishes.append({'name': name, 'description': description, 'price': price})
    wb.close()
    return dishes


def _parse_dishes_from_docx(file_obj):
    """Parse uploaded Word (.docx) file into dish dicts."""
    import docx
    doc = docx.Document(file_obj)
    raw_lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            raw_lines.append(text)
    # Also gather table rows
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                raw_lines.append(' - '.join(cells))
    return _parse_dishes_from_text('\n'.join(raw_lines))


@login_required
def owner_dish_base(request):
    """Moja Baza Potraw — per-restaurant dish library."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        return redirect("home")
    if not restaurant:
        return redirect("owner_restaurant_create")

    categories = MenuItem.Category.choices

    if request.method == "POST":
        action = request.POST.get("action")

        # ── Dodaj ręcznie ──
        if action == "add":
            cat = request.POST.get("category", "main")
            name = request.POST.get("name", "").strip()
            desc = request.POST.get("description", "").strip()
            price = request.POST.get("price", "0")
            if name:
                try:
                    price = float(str(price).replace(",", "."))
                except ValueError:
                    price = 0
                _, created = Dish.objects.get_or_create(
                    restaurant=restaurant, category=cat, name=name,
                    defaults={"description": desc, "price": price, "source": Dish.Source.MANUAL},
                )
                if created:
                    messages.success(request, f'Dodano „{name}" do bazy potraw.')
                else:
                    messages.info(request, f'„{name}" już istnieje w bazie.')

        # ── Edytuj ──
        elif action == "edit":
            dish_id = request.POST.get("dish_id")
            dish = Dish.objects.filter(id=dish_id, restaurant=restaurant).first()
            if dish:
                dish.category = request.POST.get("category", dish.category)
                dish.name = request.POST.get("name", dish.name).strip()
                dish.description = request.POST.get("description", "").strip()
                try:
                    dish.price = float(str(request.POST.get("price", dish.price)).replace(",", "."))
                except ValueError:
                    pass
                dish.save()
                messages.success(request, f'Zaktualizowano „{dish.name}".')

        # ── Usuń ──
        elif action == "delete":
            dish_id = request.POST.get("dish_id")
            dish = Dish.objects.filter(id=dish_id, restaurant=restaurant).first()
            if dish:
                name = dish.name
                dish.delete()
                messages.success(request, f'Usunięto „{name}" z bazy potraw.')

        # ── Usuń zaznaczone (bulk) ──
        elif action == "bulk_delete":
            ids = request.POST.getlist("selected_dishes")
            count = Dish.objects.filter(id__in=ids, restaurant=restaurant).delete()[0]
            messages.success(request, f'Usunięto {count} pozycji z bazy potraw.')

        # ── Import z tekstu ──
        elif action == "import_text":
            raw = request.POST.get("raw_text", "")
            cat = request.POST.get("import_category", "main")
            parsed = _parse_dishes_from_text(raw)
            added = 0
            for d in parsed:
                _, created = Dish.objects.get_or_create(
                    restaurant=restaurant, category=cat, name=d["name"],
                    defaults={"description": d["description"], "price": d["price"],
                              "source": Dish.Source.TEXT},
                )
                if created:
                    added += 1
            messages.success(request, f'Zaimportowano {added} nowych potraw z tekstu ({len(parsed)} rozpoznanych).')

        # ── Import z Excel ──
        elif action == "import_excel":
            f = request.FILES.get("excel_file")
            cat = request.POST.get("import_category", "main")
            if f:
                try:
                    parsed = _parse_dishes_from_excel(f)
                    added = 0
                    for d in parsed:
                        _, created = Dish.objects.get_or_create(
                            restaurant=restaurant, category=cat, name=d["name"],
                            defaults={"description": d["description"], "price": d["price"],
                                      "source": Dish.Source.EXCEL},
                        )
                        if created:
                            added += 1
                    messages.success(request, f'Zaimportowano {added} nowych potraw z Excel ({len(parsed)} rozpoznanych).')
                except Exception as e:
                    messages.error(request, f'Błąd importu Excel: {e}')
            else:
                messages.error(request, 'Nie wybrano pliku Excel.')

        # ── Import z Word ──
        elif action == "import_word":
            f = request.FILES.get("word_file")
            cat = request.POST.get("import_category", "main")
            if f:
                try:
                    parsed = _parse_dishes_from_docx(f)
                    added = 0
                    for d in parsed:
                        _, created = Dish.objects.get_or_create(
                            restaurant=restaurant, category=cat, name=d["name"],
                            defaults={"description": d["description"], "price": d["price"],
                                      "source": Dish.Source.WORD},
                        )
                        if created:
                            added += 1
                    messages.success(request, f'Zaimportowano {added} nowych potraw z Word ({len(parsed)} rozpoznanych).')
                except Exception as e:
                    messages.error(request, f'Błąd importu Word: {e}')
            else:
                messages.error(request, 'Nie wybrano pliku Word.')

        # ── Import z istniejącego menu ──
        elif action == "import_menu":
            menu_id = request.POST.get("menu_id")
            menu_obj = Menu.objects.filter(id=menu_id, restaurant=restaurant).first()
            if menu_obj:
                items = MenuItem.objects.filter(menu=menu_obj)
                added = 0
                for item in items:
                    _, created = Dish.objects.get_or_create(
                        restaurant=restaurant, category=item.category, name=item.name,
                        defaults={"description": item.description, "price": item.price,
                                  "source": Dish.Source.MENU},
                    )
                    if created:
                        added += 1
                messages.success(request, f'Zaimportowano {added} nowych potraw z menu „{menu_obj.name}".')

        # ── Dodaj z zewnętrznej bazy ──
        elif action == "add_external":
            ext_name = request.POST.get("name", "").strip()
            ext_cat = request.POST.get("category", "main")
            ext_price = request.POST.get("price", "0")
            if ext_name:
                try:
                    ext_price = float(str(ext_price).replace(",", "."))
                except ValueError:
                    ext_price = 0
                _, created = Dish.objects.get_or_create(
                    restaurant=restaurant, category=ext_cat, name=ext_name,
                    defaults={"description": "", "price": ext_price,
                              "source": Dish.Source.EXTERNAL},
                )
                if created:
                    messages.success(request, f'Dodano „{ext_name}" z zewnętrznej bazy.')
                else:
                    messages.info(request, f'„{ext_name}" już istnieje w Twojej bazie.')

        # ── Dodaj do menu ──
        elif action == "add_to_menu":
            dish_ids = request.POST.getlist("selected_dishes")
            menu_id = request.POST.get("target_menu_id")
            menu_obj = Menu.objects.filter(id=menu_id, restaurant=restaurant).first()
            if menu_obj and dish_ids:
                added = 0
                for dish in Dish.objects.filter(id__in=dish_ids, restaurant=restaurant):
                    if not MenuItem.objects.filter(menu=menu_obj, name=dish.name, category=dish.category).exists():
                        MenuItem.objects.create(
                            restaurant=restaurant, menu=menu_obj,
                            category=dish.category, name=dish.name,
                            description=dish.description, price=dish.price,
                        )
                        added += 1
                messages.success(request, f'Dodano {added} pozycji do menu „{menu_obj.name}".')

        return redirect("owner_dish_base")

    # ── GET — lista ──
    filter_cat = request.GET.get("category", "")
    search_q = request.GET.get("q", "").strip()

    dishes_qs = Dish.objects.filter(restaurant=restaurant)
    if filter_cat:
        dishes_qs = dishes_qs.filter(category=filter_cat)
    if search_q:
        dishes_qs = dishes_qs.filter(Q(name__icontains=search_q) | Q(description__icontains=search_q))

    # Group by category
    dish_by_category = []
    for cat_val, cat_label in categories:
        items = [d for d in dishes_qs if d.category == cat_val]
        if items or not filter_cat:
            dish_by_category.append({
                "key": cat_val,
                "label": cat_label,
                "items": items,
                "count": len(items),
            })

    all_menus = Menu.objects.filter(restaurant=restaurant)

    return render(request, "bookings/owner/dish_base.html", {
        "restaurant": restaurant,
        "categories": categories,
        "dish_by_category": dish_by_category,
        "total_dishes": dishes_qs.count(),
        "filter_cat": filter_cat,
        "search_q": search_q,
        "all_menus": all_menus,
        "source_labels": DISH_SOURCE_LABELS,
        "source_icons": DISH_SOURCE_ICONS,
    })


@login_required
def dish_base_api(request):
    """AJAX API — search dishes in own library + external catalog."""
    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None:
        return JsonResponse([], safe=False)

    q = request.GET.get("q", "").strip()
    category = request.GET.get("category", "").strip()
    source = request.GET.get("source", "all")  # "own", "external", "all"

    results = []

    # Own dish base
    if source in ("own", "all"):
        qs = Dish.objects.filter(restaurant=restaurant)
        if category:
            qs = qs.filter(category=category)
        if q:
            qs = qs.filter(name__icontains=q)
        for d in qs[:20]:
            results.append({
                "id": d.id,
                "name": d.name,
                "category": d.category,
                "description": d.description,
                "price": float(d.price),
                "source": "own",
            })

    # External catalog (MenuItemTemplate)
    if source in ("external", "all"):
        qs = MenuItemTemplate.objects.all()
        if category:
            qs = qs.filter(category=category)
        if q:
            qs = qs.filter(name__icontains=q)
        own_names = set(Dish.objects.filter(restaurant=restaurant).values_list("name", flat=True))
        for t in qs[:20]:
            if t.name not in own_names:
                results.append({
                    "id": None,
                    "name": t.name,
                    "category": t.category,
                    "description": "",
                    "price": float(t.last_price),
                    "source": "external",
                })

    return JsonResponse(results[:30], safe=False)


# ── API podpowiedzi menu ──────────────────────────────────────────────────────

@login_required
def menu_suggestions_api(request):
    """Zwraca JSON z podpowiedziami pozycji menu (autocomplete).

    Priorytet: 1) Baza potraw właściciela, 2) Globalny katalog MenuItemTemplate.
    """
    q = request.GET.get("q", "").strip()
    category = request.GET.get("category", "").strip()

    results = []
    seen_names = set()

    # 1) Baza potraw restauracji (jeśli owner jest zalogowany)
    memberships, restaurant, membership = _get_owner_context(request)
    if restaurant:
        dish_qs = Dish.objects.filter(restaurant=restaurant)
        if category:
            dish_qs = dish_qs.filter(category=category)
        if q:
            dish_qs = dish_qs.filter(name__icontains=q)
        for d in dish_qs[:15]:
            results.append({
                "name": d.name,
                "category": d.category,
                "last_price": float(d.price),
                "source": "dish_base",
            })
            seen_names.add(d.name.lower())

    # 2) Globalny katalog
    tpl_qs = MenuItemTemplate.objects.all()
    if category:
        tpl_qs = tpl_qs.filter(category=category)
    if q:
        tpl_qs = tpl_qs.filter(name__icontains=q)
    for t in tpl_qs[:20]:
        if t.name.lower() not in seen_names:
            results.append({
                "name": t.name,
                "category": t.category,
                "last_price": float(t.last_price),
                "source": "template",
            })
            seen_names.add(t.name.lower())

    return JsonResponse(results[:25], safe=False)


# ── Zapisane menu ──────────────────────────────────────────────────────────────

@login_required
def toggle_save_menu(request, restaurant_pk):
    """Zapisz / usuń menu restauracji z ulubionych."""
    restaurant = get_object_or_404(Restaurant, pk=restaurant_pk, is_active=True)
    saved, created = SavedMenu.objects.get_or_create(
        user=request.user, restaurant=restaurant,
    )
    if not created:
        saved.delete()
        messages.info(request, "Usunięto menu z zapisanych.")
    else:
        messages.success(request, "Zapisano menu!")
    next_url = request.GET.get("next")
    if next_url:
        return redirect(next_url)
    return redirect("restaurant_detail", pk=restaurant_pk)


@login_required
def saved_menus_list(request):
    """Lista zapisanych menu użytkownika."""
    saved = SavedMenu.objects.filter(user=request.user).select_related(
        "restaurant"
    )
    # Wczytaj menu dla każdej firmy
    entries = []
    for s in saved:
        visible_menu = MenuItem.objects.filter(
            restaurant=s.restaurant, is_visible=True
        )
        categories = MenuItem.Category.choices
        menu_by_category = []
        for cat_value, cat_label in categories:
            items = visible_menu.filter(category=cat_value)
            if items.exists():
                menu_by_category.append({"label": cat_label, "items": items})
        entries.append({
            "saved": s,
            "restaurant": s.restaurant,
            "menu_by_category": menu_by_category,
        })
    return render(request, "bookings/saved_menus.html", {
        "entries": entries,
    })


# ── Formularz rezerwacji na zewnętrzną stronę (embed) ──────────────────────────

def _build_calendar(restaurant, cal_year, cal_month):
    """Helper: generate calendar data for a given month (legacy wrapper)."""
    cal = build_month_grid(restaurant, cal_year, cal_month)
    return cal["weeks"], cal["year"], cal["month"], cal["prev"], cal["next"]


def embed_booking(request, slug):
    """Publiczny formularz rezerwacji – strona pod /rezerwacja/<slug>/."""
    from django.contrib.auth.models import User as AuthUser

    restaurant = get_object_or_404(
        Restaurant, booking_slug=slug, embed_enabled=True, is_active=True,
    )
    is_attraction = restaurant.firm_type == Restaurant.FirmType.ATTRACTION

    # Kalendarz
    today = timezone.now().date()
    try:
        cal_year = int(request.GET.get("cal_year", today.year))
        cal_month = int(request.GET.get("cal_month", today.month))
    except (ValueError, TypeError):
        cal_year, cal_month = today.year, today.month

    show_cal = restaurant.firm_type == Restaurant.FirmType.VENUE and restaurant.show_calendar
    cal = None
    if show_cal:
        cal = build_month_grid(restaurant, cal_year, cal_month)

    ctx = {
        "restaurant": restaurant,
        "is_attraction": is_attraction,
        "cal": cal,
    }

    # Sprawdź czy firma ma menu (do wyświetlenia steppera)
    _active_menu = Menu.objects.filter(restaurant=restaurant, is_active=True).first()
    ctx["has_menu"] = _active_menu and MenuItem.objects.filter(
        restaurant=restaurant, is_visible=True, menu=_active_menu,
    ).exists()

    if request.method == "POST":
        form = BookingForm(request.POST, firm_type=restaurant.firm_type)
        if form.is_valid():
            booking = form.save(commit=False)
            booking.restaurant = restaurant

            guest_ok = True
            if not is_attraction and booking.guest_count and booking.guest_count > restaurant.max_guests:
                form.add_error(
                    "guest_count",
                    f"Firma przyjmuje maksymalnie {restaurant.max_guests} gości.",
                )
                guest_ok = False

            if guest_ok:
                if Booking.objects.filter(
                    restaurant=restaurant,
                    event_date=booking.event_date,
                ).exclude(status=Booking.Status.CANCELLED).exists():
                    form.add_error("event_date", "Ta data jest już zarezerwowana.")
                else:
                    # Przypisz lub utwórz użytkownika-gościa na podstawie e-mail
                    email = form.cleaned_data["email"]
                    guest_user, created = AuthUser.objects.get_or_create(
                        email=email,
                        defaults={
                            "username": f"guest_{email.split('@')[0]}_{AuthUser.objects.count()}",
                            "first_name": form.cleaned_data.get("first_name", ""),
                            "last_name": form.cleaned_data.get("last_name", ""),
                        },
                    )
                    if created:
                        guest_user.set_unusable_password()
                        guest_user.save()

                    booking.user = guest_user
                    booking.save()

                    # Wiadomość powitalna
                    if restaurant.welcome_message.strip():
                        owner_qs = restaurant.owners.filter(role="owner").select_related("user")
                        if not owner_qs.exists():
                            owner_qs = restaurant.owners.select_related("user")
                        if owner_qs.exists():
                            BookingMessage.objects.create(
                                booking=booking,
                                sender=owner_qs.first().user,
                                content=restaurant.welcome_message.strip(),
                            )

                    # Sprawdź czy firma ma aktywne menu → krok 2
                    active_menu = Menu.objects.filter(restaurant=restaurant, is_active=True).first()
                    has_menu = active_menu and MenuItem.objects.filter(
                        restaurant=restaurant, is_visible=True, menu=active_menu,
                    ).exists()
                    if has_menu:
                        return redirect("embed_booking_menu", slug=slug, booking_pk=booking.pk)

                    ctx["success"] = True
                    ctx["booking"] = booking
                    ctx["form"] = form
                    return render(request, "bookings/embed_booking.html", ctx)
    else:
        form = BookingForm(firm_type=restaurant.firm_type)

    ctx["form"] = form
    return render(request, "bookings/embed_booking.html", ctx)


def embed_calendar_partial(request, slug):
    """AJAX – kalendarz dla embed formularza."""
    restaurant = get_object_or_404(
        Restaurant, booking_slug=slug, embed_enabled=True, is_active=True,
    )
    cal = build_month_grid(
        restaurant,
        request.GET.get("cal_year"),
        request.GET.get("cal_month"),
    )
    return render(request, "bookings/_embed_calendar_partial.html", {"cal": cal})


def embed_booking_menu(request, slug, booking_pk):
    """Krok 2: wybór menu w publicznym formularzu rezerwacji."""
    restaurant = get_object_or_404(
        Restaurant, booking_slug=slug, embed_enabled=True, is_active=True,
    )
    booking = get_object_or_404(Booking, pk=booking_pk, restaurant=restaurant)

    active_menu = Menu.objects.filter(restaurant=restaurant, is_active=True).first()
    menu_items = MenuItem.objects.filter(
        restaurant=restaurant, is_visible=True, menu=active_menu,
    ) if active_menu else MenuItem.objects.none()

    if not menu_items.exists():
        return render(request, "bookings/embed_booking.html", {
            "restaurant": restaurant,
            "success": True,
            "booking": booking,
        })

    # Ustaw typ menu jeśli brak
    if not booking.menu_type:
        enabled = restaurant.enabled_menu_types or ["detailed"]
        booking.menu_type = enabled[0] if len(enabled) == 1 else "detailed"
        booking.save(update_fields=["menu_type"])

    if request.method == "POST":
        # Zapisz wybrane pozycje menu
        booking.menu_selections.all().delete()
        for item in menu_items:
            qty_str = request.POST.get(f"qty_{item.id}", "0")
            try:
                qty = int(qty_str)
            except ValueError:
                qty = 0
            if qty > 0:
                BookingMenuItem.objects.create(
                    booking=booking,
                    menu_item=item,
                    quantity=qty,
                )

        return render(request, "bookings/embed_booking_menu.html", {
            "restaurant": restaurant,
            "booking": booking,
            "success": True,
        })

    # Przygotuj dane menu
    menu_by_category = _build_menu_by_category(menu_items, booking)

    return render(request, "bookings/embed_booking_menu.html", {
        "restaurant": restaurant,
        "booking": booking,
        "menu_by_category": menu_by_category,
    })


def embed_booking_menu_skip(request, slug, booking_pk):
    """Pominięcie wyboru menu — przejdź do ekranu sukcesu."""
    restaurant = get_object_or_404(
        Restaurant, booking_slug=slug, embed_enabled=True, is_active=True,
    )
    booking = get_object_or_404(Booking, pk=booking_pk, restaurant=restaurant)
    return render(request, "bookings/embed_booking_menu.html", {
        "restaurant": restaurant,
        "booking": booking,
        "success": True,
    })


@login_required
def owner_generate_embed(request):
    """Generuje (lub regeneruje) formularz rezerwacji dla firmy."""
    from django.utils.text import slugify

    memberships, restaurant, membership = _get_owner_context(request)
    if memberships is None or not restaurant:
        messages.error(request, "Nie masz firmy do skonfigurowania.")
        return redirect("owner_dashboard")

    if request.method == "POST":
        action = request.POST.get("action")

        if action == "generate":
            # Generuj slug jeśli brak
            if not restaurant.booking_slug:
                base_slug = slugify(restaurant.name)
                slug = base_slug
                counter = 1
                while Restaurant.objects.filter(booking_slug=slug).exclude(pk=restaurant.pk).exists():
                    slug = f"{base_slug}-{counter}"
                    counter += 1
                restaurant.booking_slug = slug

            # Scrapuj style ze strony www
            if restaurant.website:
                try:
                    css = scrape_styles(restaurant.website)
                    restaurant.scraped_css = css
                except Exception:
                    restaurant.scraped_css = ""

            restaurant.embed_enabled = True
            restaurant.save()
            messages.success(request, "Formularz rezerwacji został wygenerowany!")

        elif action == "disable":
            restaurant.embed_enabled = False
            restaurant.save()
            messages.info(request, "Formularz zewnętrzny został wyłączony.")

        elif action == "refresh_css":
            if restaurant.website:
                try:
                    css = scrape_styles(restaurant.website)
                    restaurant.scraped_css = css
                    restaurant.save()
                    messages.success(request, "Style zostały odświeżone ze strony www.")
                except Exception:
                    messages.warning(request, "Nie udało się pobrać stylów ze strony.")
            else:
                messages.warning(request, "Brak adresu strony www w ustawieniach firmy.")

    embed_url = None
    if restaurant.embed_enabled and restaurant.booking_slug:
        embed_url = request.build_absolute_uri(f"/rezerwacja/{restaurant.booking_slug}/")

    return render(request, "bookings/owner/embed_settings.html", {
        "restaurant": restaurant,
        "embed_url": embed_url,
    })
